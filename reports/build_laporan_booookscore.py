from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "reports/Laporan_Proyek_Tahap_1_BooookScore.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9DEE7"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color="2E74B5", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_body_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.style = "Normal"
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    p.style = f"Heading {level}"
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=font_size, color=INK, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    set_run_font(r, size=font_size, color=RGBColor(30, 30, 30))
    set_table_geometry(table, widths)
    set_table_borders(table)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for r in p2.runs:
        set_run_font(r, size=10.5)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="C8D3E2")
    doc.add_paragraph()


def add_numbered_source(doc, number, label, url):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(f"{number}. {label}: ")
    set_run_font(r, size=10)
    add_hyperlink(p, url, url)


def apply_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 22, INK, 0, 6),
        ("Subtitle", 12, GRAY, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name in ("Title", "Heading 1", "Heading 2", "Heading 3")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def build_doc():
    doc = Document()
    apply_styles(doc)
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.text = "Project Akhir Individu DSS | Summarization"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = "Laporan Proyek Tahap 1 - BooookScore"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=GRAY)

    p = doc.add_paragraph("LAPORAN PROYEK TAHAP 1")
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Project Akhir Individu Decision Support System")
    p.style = "Subtitle"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Task NLP: Summarization")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)

    meta = [
        ("Nama Repository", "BooookScore"),
        ("URL Repository", "https://github.com/lilakk/BooookScore"),
        ("Paper Rujukan Utama", "https://arxiv.org/pdf/2310.00785.pdf"),
        ("Deadline Kelas B", "Senin, 11 Mei 2026, 23.59 WIB"),
    ]
    add_table(doc, ["Komponen", "Keterangan"], meta, [2600, 6760], BLUE_GRAY, font_size=10)
    paragraph_border_bottom(doc.add_paragraph(), color="2E74B5", size="10")

    add_heading(doc, "1. Pendahuluan", 1)
    add_body_para(doc, "Summarization merupakan salah satu task penting dalam Natural Language Processing (NLP) yang bertujuan menghasilkan ringkasan dari teks panjang dengan tetap mempertahankan informasi utama. Task ini banyak digunakan untuk membantu pengguna memahami isi dokumen secara cepat, terutama ketika dokumen memiliki ukuran besar, seperti artikel ilmiah, laporan bisnis, berita, dokumen hukum, transkrip rapat, hingga buku.")
    add_body_para(doc, "Dalam konteks Decision Support System (DSS), summarization memiliki peran penting karena sistem pendukung keputusan sering berhadapan dengan informasi yang panjang, kompleks, dan tidak terstruktur. Dengan adanya summarization, pengambil keputusan dapat memperoleh inti informasi secara lebih cepat, sehingga proses analisis, evaluasi alternatif, dan pengambilan keputusan dapat dilakukan dengan lebih efisien.")
    add_body_para(doc, "Repository yang dipilih dalam Project Tahap 1 ini adalah BooookScore. Repository ini dipilih karena secara langsung membahas task book-length summarization, yaitu peringkasan teks sangat panjang dalam bentuk buku. BooookScore relevan dengan perkembangan NLP modern karena menggunakan Large Language Models (LLMs) dan prompt methods untuk menghasilkan serta mengevaluasi ringkasan.")
    add_callout(doc, "Alasan utama pemilihan repository", "BooookScore sesuai dengan kebutuhan tugas karena menyediakan repository GitHub, code, contoh data, prompt, hasil ringkasan, serta paper rujukan utama yang jelas dan saling berhubungan.")

    add_heading(doc, "2. Nama URL/GitHub dan Nama Pemilik", 1)
    add_table(doc, ["Komponen", "Keterangan"], [
        ("Nama Repository", "BooookScore"),
        ("URL GitHub", "https://github.com/lilakk/BooookScore"),
        ("Pemilik Repository", "lilakk / Yapei Chang dan tim"),
        ("Jenis Repository", "Official code and data release"),
        ("Keterkaitan Paper", "Repository resmi dari paper BooookScore ICLR 2024"),
    ], [2600, 6760], BLUE_GRAY, font_size=10)
    add_body_para(doc, "Repository BooookScore merupakan repository resmi yang menyediakan code dan data release untuk paper BooookScore. Pada halaman GitHub, repository ini dijelaskan sebagai package untuk menghasilkan ringkasan teks panjang dan mengevaluasi koherensi ringkasan tersebut.")

    add_heading(doc, "3. Spesifikasi Paper Rujukan Utama", 1)
    add_table(doc, ["Komponen", "Keterangan"], [
        ("Judul Paper", "BooookScore: A systematic exploration of book-length summarization in the era of LLMs"),
        ("Penulis", "Yapei Chang, Kyle Lo, Tanya Goyal, Mohit Iyyer"),
        ("Tahun Terbit", "2024"),
        ("Conference/Jurnal", "The Twelfth International Conference on Learning Representations / ICLR 2024"),
        ("URL Paper", "https://arxiv.org/pdf/2310.00785.pdf"),
        ("Task NLP", "Summarization, khususnya book-length summarization"),
        ("Metode Utama", "Large Language Models, prompt-based summarization, hierarchical merging, incremental updating, automatic evaluation metric BooookScore"),
        ("Hubungan Paper dengan Repository", "Repository GitHub BooookScore merupakan official code and data release dari paper tersebut"),
    ], [2600, 6760], BLUE_GRAY, font_size=9.2)
    add_body_para(doc, "Paper ini membahas eksplorasi sistematis terhadap summarization untuk teks sangat panjang, khususnya teks berbentuk buku. Permasalahan utama yang diangkat adalah bahwa buku dapat memiliki panjang lebih dari 100.000 token, sehingga tidak selalu dapat langsung dimasukkan ke dalam context window LLM.")
    add_body_para(doc, "Paper ini membandingkan dua workflow utama berbasis prompt, yaitu hierarchical merging dan incremental updating. Selain menghasilkan ringkasan, paper ini juga memperkenalkan BooookScore sebagai metrik evaluasi otomatis untuk mengukur kualitas koherensi ringkasan panjang yang dihasilkan oleh LLM.")

    add_heading(doc, "4. Spesifikasi Code", 1)
    add_body_para(doc, "Repository BooookScore menggunakan bahasa pemrograman utama Python. Code dalam repository ini disusun sebagai package yang dapat digunakan melalui perintah Python module, misalnya python -m booookscore.chunk, python -m booookscore.summ, python -m booookscore.postprocess, dan python -m booookscore.score.")
    add_table(doc, ["No", "Komponen", "Keterangan"], [
        ("1", "Bahasa Pemrograman Utama", "Python"),
        ("2", "Jenis File yang Tersedia", "File Python package, file prompt, file pickle dataset, file JSON summaries, README, LICENSE, requirements.txt"),
        ("3", "Folder Utama Repository", "annotations, booookscore, data, misc, prompts, summaries"),
        ("4", "Folder Code Utama", "booookscore"),
        ("5", "Folder Prompt", "prompts"),
        ("6", "Folder Data", "data"),
        ("7", "Folder Summaries", "summaries"),
        ("8", "File Requirements", "requirements.txt"),
        ("9", "Lisensi Repository", "MIT License"),
        ("10", "Deskripsi Umum Fungsi Code", "Package untuk menghasilkan ringkasan teks panjang dan mengevaluasi koherensi ringkasan menggunakan BooookScore"),
        ("11", "API LLM yang Didukung", "OpenAI, Anthropic, Together"),
        ("12", "Metode Summarization", "Hierarchical merging dan incremental updating"),
        ("13", "Format Input Utama", "Pickle dictionary berisi nama buku dan teks lengkap buku"),
        ("14", "Format Output Utama", "JSON berisi nama buku dan hasil ringkasan akhir"),
    ], [650, 2550, 6160], BLUE_GRAY, font_size=8.8)
    add_body_para(doc, "Secara umum, code BooookScore digunakan untuk menyiapkan data buku, melakukan chunking teks panjang, menjalankan summarization menggunakan LLM, menggunakan prompt untuk menghasilkan ringkasan, menggunakan metode hierarchical merging dan incremental updating, melakukan post-processing hasil summary, serta menyimpan hasil ringkasan dalam format JSON.")

    doc.add_page_break()
    add_heading(doc, "5. Spesifikasi Dataset", 1)
    add_body_para(doc, "Dataset pada repository BooookScore digunakan untuk task book-length summarization, yaitu peringkasan teks buku penuh. Dataset berbentuk file pickle dictionary dengan ekstensi .pkl. Struktur datanya adalah key = nama buku dan value = isi teks lengkap buku. Contoh file dataset yang tersedia dalam repository adalah data/example_all_books.pkl.")
    add_table(doc, ["No", "Komponen", "Keterangan"], [
        ("1", "Jenis Dataset", "Teks buku untuk book-length summarization"),
        ("2", "Format Dataset", ".pkl / pickle dictionary"),
        ("3", "Struktur Dataset", "Key berupa nama buku, value berupa isi teks lengkap buku"),
        ("4", "Contoh File Dataset", "data/example_all_books.pkl"),
        ("5", "Unit Data", "Satu data merepresentasikan satu buku penuh"),
        ("6", "Input Proses", "Teks lengkap buku"),
        ("7", "Proses Awal", "Chunking teks panjang sesuai ukuran chunk"),
        ("8", "Output Summary", "File JSON"),
        ("9", "Folder Output Summary", "summaries"),
        ("10", "Isi Folder Summaries", "Hasil ringkasan dari berbagai model LLM dan metode summarization"),
    ], [650, 2550, 6160], BLUE_GRAY, font_size=9)
    add_body_para(doc, "Folder summaries berisi file JSON hasil ringkasan dari berbagai model dan konfigurasi, misalnya hasil ringkasan menggunakan metode hierarchical maupun incremental. File yang berakhiran cleaned.json berisi final summaries yang sudah melalui post-processing, sedangkan file yang tidak berakhiran cleaned.json berisi intermediate summaries.")

    add_heading(doc, "6. Sampling Isi Dataset", 1)
    add_body_para(doc, "Karena dataset BooookScore berbentuk .pkl, isi dataset tidak dapat langsung dibaca seperti file CSV atau Excel. Dataset perlu dibuka menggunakan Python, misalnya dengan library pickle. Secara konseptual, dataset berbentuk dictionary dengan nama buku sebagai key dan isi teks lengkap buku sebagai value.")
    add_table(doc, ["No", "Key / Nama Buku", "Value / Isi Data", "Penjelasan"], [
        ("1", "Judul Buku A", "Teks lengkap buku A", "Digunakan sebagai input utama untuk proses summarization"),
        ("2", "Judul Buku B", "Teks lengkap buku B", "Dipotong menjadi beberapa chunk agar dapat diproses oleh LLM"),
        ("3", "Judul Buku C", "Teks lengkap buku C", "Diproses hingga menghasilkan summary akhir dalam format JSON"),
    ], [650, 2100, 2500, 4110], BLUE_GRAY, font_size=9)
    add_body_para(doc, "Setiap data merepresentasikan satu buku penuh. Nama buku menjadi identifier, sedangkan isi teks buku menjadi input untuk proses summarization. Karena teks buku memiliki ukuran sangat panjang, teks tersebut diproses melalui tahap chunking agar setiap bagian dapat masuk ke context window LLM.")

    add_heading(doc, "7. Teknik / Cara Kerja", 1)
    add_body_para(doc, "BooookScore bekerja dengan alur utama yang dirancang untuk menangani teks sangat panjang. Input awal berupa file pickle yang berisi dictionary nama buku dan teks lengkap buku. Karena teks buku dapat melebihi context window LLM, teks tersebut terlebih dahulu dipotong menjadi beberapa chunk. Setiap chunk kemudian diberikan kepada LLM menggunakan prompt tertentu agar model dapat menghasilkan ringkasan bagian.")
    add_body_para(doc, "Setelah ringkasan per chunk diperoleh, BooookScore menggunakan dua metode utama. Metode pertama adalah hierarchical merging, yaitu menggabungkan ringkasan-ringkasan kecil secara bertingkat hingga menjadi satu ringkasan akhir. Metode kedua adalah incremental updating, yaitu memperbarui ringkasan secara bertahap ketika chunk baru diproses. Hasil akhir kemudian disimpan dalam format JSON dan dapat dievaluasi menggunakan metrik BooookScore.")
    add_table(doc, ["No", "Tahap", "Penjelasan"], [
        ("1", "Input Data", "Data berupa teks buku dalam file pickle dictionary"),
        ("2", "Chunking", "Teks panjang dipotong menjadi beberapa chunk"),
        ("3", "Prompting", "Prompt digunakan untuk meminta LLM menghasilkan ringkasan"),
        ("4", "Summary per Chunk", "LLM menghasilkan ringkasan untuk setiap bagian teks"),
        ("5", "Hierarchical Merging", "Ringkasan kecil digabungkan secara bertingkat menjadi ringkasan akhir"),
        ("6", "Incremental Updating", "Ringkasan diperbarui secara bertahap seiring masuknya chunk baru"),
        ("7", "Post-processing", "Hasil ringkasan dibersihkan dari artifact tertentu"),
        ("8", "Penyimpanan Output", "Summary akhir disimpan dalam format JSON"),
        ("9", "Evaluasi", "Summary dapat dievaluasi menggunakan BooookScore"),
    ], [650, 2500, 6210], BLUE_GRAY, font_size=9)

    doc.add_page_break()
    add_heading(doc, "8. Kesesuaian dengan Requirements Tugas", 1)
    add_table(doc, ["No", "Requirement Tugas", "Status", "Keterangan"], [
        ("1", "Memiliki GitHub/repository", "Sesuai", "Repository tersedia di GitHub: https://github.com/lilakk/BooookScore"),
        ("2", "Memiliki code", "Sesuai", "Terdapat package Python dalam folder booookscore"),
        ("3", "Memiliki dataset atau contoh data", "Sesuai", "Terdapat contoh data data/example_all_books.pkl"),
        ("4", "Memiliki paper rujukan utama", "Sesuai", "Paper ICLR 2024 tersedia di arXiv"),
        ("5", "Paper dan repository saling berhubungan", "Sesuai", "Repository merupakan official code and data release dari paper BooookScore"),
        ("6", "Task sesuai dengan Summarization", "Sesuai", "Fokus utama adalah book-length summarization"),
        ("7", "Memakai LLM/prompt methods", "Sesuai", "Menggunakan LLM dan prompt workflow"),
        ("8", "Rilis dalam 3 tahun terakhir", "Sesuai", "Paper tahun 2024 dan repository aktif untuk paper ICLR 2024"),
        ("9", "Tidak wajib memakai API berbayar pada tahap 1", "Sesuai", "Tahap 1 hanya memilih dan menganalisis repository, belum menjalankan program"),
        ("10", "Bisa dipelajari untuk tahap implementasi selanjutnya", "Sesuai", "Repository memiliki README, requirements, contoh data, prompt, code, dan output summaries"),
    ], [550, 2750, 1250, 4810], BLUE_GRAY, font_size=8.2)
    add_body_para(doc, "Berdasarkan checklist di atas, repository BooookScore memenuhi requirements Project Tahap 1. Repository ini memiliki hubungan langsung dengan paper rujukan utama, memiliki code, prompt, contoh dataset, serta hasil summary yang dapat digunakan sebagai bahan pembelajaran untuk tahap implementasi berikutnya.")

    add_heading(doc, "9. Kesimpulan", 1)
    add_body_para(doc, "BooookScore layak digunakan sebagai repository utama untuk Project Akhir Individu DSS Tahap 1 dengan task Summarization. Repository ini sesuai karena membahas book-length summarization menggunakan Large Language Models dan prompt methods. Selain itu, BooookScore memiliki paper rujukan utama yang jelas, yaitu paper ICLR 2024 berjudul BooookScore: A systematic exploration of book-length summarization in the era of LLMs.")
    add_body_para(doc, "Repository ini juga menyediakan code, contoh dataset, prompt, file requirements, lisensi, serta hasil summary dalam format JSON. Alur kerjanya mencakup chunking teks panjang, summarization menggunakan LLM, hierarchical merging, incremental updating, post-processing, dan evaluasi summary menggunakan BooookScore. Dengan demikian, BooookScore merupakan pilihan yang tepat dan relevan untuk dilanjutkan pada tahap implementasi berikutnya.")

    doc.add_page_break()
    add_heading(doc, "Daftar Sumber", 1)
    add_numbered_source(doc, 1, "Repository BooookScore", "https://github.com/lilakk/BooookScore")
    add_numbered_source(doc, 2, "Paper BooookScore ICLR 2024", "https://arxiv.org/pdf/2310.00785.pdf")
    add_numbered_source(doc, 3, "Folder Data BooookScore", "https://github.com/lilakk/BooookScore/tree/main/data")
    add_numbered_source(doc, 4, "Folder Prompts BooookScore", "https://github.com/lilakk/BooookScore/tree/main/prompts")
    add_numbered_source(doc, 5, "Folder Summaries BooookScore", "https://github.com/lilakk/BooookScore/tree/main/summaries")
    add_numbered_source(doc, 6, "Requirements BooookScore", "https://github.com/lilakk/BooookScore/blob/main/requirements.txt")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
